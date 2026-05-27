"""Build a one-page proforma for the JK Tyre BTP V1 Forward Scheduler.

Mirrors the layout / palette / typography of
  docs/Agentic_Workflow_Proforma_Faizan 1.docx
but with content adapted to the BTP scheduler.

Outputs:
  docs/BTP_V1_Proforma_AnmolSaini.docx
  docs/BTP_V1_Proforma_AnmolSaini.pdf   (reportlab rendition of the same layout)

Run:  uv run python docs/build_proforma.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Mm

# -----------------------------------------------------------------------------
# Palette (verbatim from the sample)
# -----------------------------------------------------------------------------
KICKER_GREEN  = "1D9E75"
HEADING_DARK  = "3A3A3A"
SUBHEAD_GREY  = "555555"
LABEL_GREY    = "888888"
ACCENT_BLUE   = "185FA5"
ACCENT_PURPLE = "534AB7"
ACCENT_AMBER  = "854F0B"
BORDER_GREY   = "CCCCCC"

CARD_BLUE     = "E6F1FB"
CARD_MINT     = "E8F7F2"
CARD_LAVENDER = "EEEDFE"
CARD_CREAM    = "F5F5F2"
CARD_PEACH    = "FAEEDA"
CARD_ROSE     = "FDEAEA"

FONT_NAME = "Arial"


# -----------------------------------------------------------------------------
# OXML helpers
# -----------------------------------------------------------------------------
def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color="FFFFFF", size=0, sides=("top", "left", "bottom", "right")) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        if side in sides:
            edge.set(qn("w:val"), "single" if size else "none")
            edge.set(qn("w:sz"), str(size))
            edge.set(qn("w:color"), color)
        else:
            edge.set(qn("w:val"), "nil")
        borders.append(edge)
    tcPr.append(borders)


def _set_cell_margins(cell, top=80, left=120, bottom=80, right=120) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _strip_table_borders(table) -> None:
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    tblPr.append(borders)


def _run(paragraph, text, *, bold=False, italic=False, color=HEADING_DARK,
         size_pt=10.0, font=FONT_NAME, spacing=None):
    r = paragraph.add_run(text)
    r.font.name = font
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for k in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{k}"), font)
    rPr.append(rFonts)
    r.font.size = Pt(size_pt)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    if spacing is not None:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rPr.append(sp)
    return r


def _para(cell_or_doc, *, space_after=2, space_before=0, align=None):
    if hasattr(cell_or_doc, "add_paragraph"):
        p = cell_or_doc.add_paragraph()
    else:
        p = cell_or_doc.paragraphs[-1]
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def _cell_first_para(cell):
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    return p


def _kicker(paragraph, text, color=KICKER_GREEN, size_pt=8.0):
    _run(paragraph, text.upper(), bold=True, color=color, size_pt=size_pt, spacing=140)


def _title(paragraph, text, color=HEADING_DARK, size_pt=21.0):
    _run(paragraph, text, bold=True, color=color, size_pt=size_pt)


def _subtitle(paragraph, text, color=SUBHEAD_GREY, size_pt=10.0):
    _run(paragraph, text, italic=True, color=color, size_pt=size_pt)


def _label_value(paragraph, label, value, value_bold=True, value_color=HEADING_DARK):
    _run(paragraph, f"{label}  ", color=LABEL_GREY, size_pt=8.5)
    _run(paragraph, value, bold=value_bold, color=value_color, size_pt=8.5)


# -----------------------------------------------------------------------------
# Page setup
# -----------------------------------------------------------------------------
def _set_margins(section, top=Cm(1.4), bottom=Cm(1.4), left=Cm(1.5), right=Cm(1.5)):
    section.top_margin = top
    section.bottom_margin = bottom
    section.left_margin = left
    section.right_margin = right


def _set_footer(section, text="CONFIDENTIAL — Internal Use Only"):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(p, text, color=LABEL_GREY, size_pt=8.0)
    _run(p, "          Page ", color=LABEL_GREY, size_pt=8.0)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    p._p.append(fld)


# -----------------------------------------------------------------------------
# Section builders
# -----------------------------------------------------------------------------
def _section_kicker(doc, label, color=ACCENT_BLUE):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    _run(p, label.upper(), bold=True, color=color, size_pt=8.5, spacing=160)


def _two_col_title_table(doc):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(12.3)
    tbl.columns[1].width = Cm(5.2)
    _strip_table_borders(tbl)
    left, right = tbl.rows[0].cells
    left.width = Cm(12.3)
    right.width = Cm(5.2)
    for c in (left, right):
        _set_cell_borders(c, sides=())
        _set_cell_margins(c, top=0, left=0, bottom=0, right=60)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # LEFT — kicker + title + subtitle
    p1 = _cell_first_para(left)
    _kicker(p1, "FORWARD PRODUCTION SCHEDULER", color=KICKER_GREEN)
    p2 = left.add_paragraph()
    _title(p2, "JK Tyre BTP — Passenger-Car Radial Pilot")
    p2.paragraph_format.space_after = Pt(2)
    p3 = left.add_paragraph()
    _subtitle(p3,
              "A deterministic, BOM-driven single-pass scheduler that converts a fixed "
              "May curing plan into a fully time-stamped, machine-assigned upstream "
              "schedule across the eight in-scope PCR components.")
    p3.paragraph_format.space_after = Pt(0)

    # RIGHT — metadata
    _set_cell_borders(right, color=BORDER_GREY, size=4, sides=("left",))
    _set_cell_margins(right, top=0, left=180, bottom=0, right=0)
    p = _cell_first_para(right)
    _label_value(p, "Author", "Anmol Saini")
    p = right.add_paragraph()
    _label_value(p, "Role", "Data Science Intern")
    p = right.add_paragraph()
    _label_value(p, "Date", "24 May 2026")
    p = right.add_paragraph()
    _label_value(p, "Team", "Data Science Team")
    p = right.add_paragraph()
    _run(p, "Readiness", color=LABEL_GREY, size_pt=8.5)
    p = right.add_paragraph()
    _run(p, "○ Concept   ", color=LABEL_GREY, size_pt=8.5)
    _run(p, "● Prototype", bold=True, color=ACCENT_BLUE, size_pt=8.5)
    _run(p, "   ○ Production", color=LABEL_GREY, size_pt=8.5)
    p = right.add_paragraph()
    _run(p, "Scope  ", color=LABEL_GREY, size_pt=8.5)
    _run(p, "V1 — Demand Fulfilment", bold=True, color=HEADING_DARK, size_pt=8.5)
    for para in right.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(3)


def _card(doc, fill_hex, accent_color, title_text, bullets, *, bullet_size=9.5):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _strip_table_borders(tbl)
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(17.5)
    _set_cell_shading(cell, fill_hex)
    _set_cell_borders(cell, sides=())
    _set_cell_margins(cell, top=180, left=240, bottom=180, right=240)
    p = _cell_first_para(cell)
    _run(p, title_text, bold=True, color=accent_color, size_pt=11.0)
    p.paragraph_format.space_after = Pt(4)
    for b in bullets:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(3)
        bp.paragraph_format.left_indent = Cm(0.35)
        _run(bp, "•  ", bold=True, color=accent_color, size_pt=bullet_size)
        # Allow markdown-ish bold via the **…** convention.
        _markdown_run(bp, b, base_color=HEADING_DARK, size_pt=bullet_size)
    return tbl


def _markdown_run(paragraph, text, *, base_color=HEADING_DARK, size_pt=9.5):
    """Tiny inline parser supporting **bold** segments."""
    import re
    pieces = re.split(r"(\*\*.+?\*\*)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            _run(paragraph, piece[2:-2], bold=True, color=base_color, size_pt=size_pt)
        else:
            _run(paragraph, piece, bold=False, color=base_color, size_pt=size_pt)


def _value_chain(doc):
    stages = [
        ("Mixing",                "Full",    CARD_MINT,     KICKER_GREEN),
        ("Final Mixing",          "Full",    CARD_MINT,     KICKER_GREEN),
        ("Calendering (FRC)",     "Full",    CARD_MINT,     KICKER_GREEN),
        ("Belt + Ply Cutter",     "Full",    CARD_MINT,     KICKER_GREEN),
        ("Extrusion",             "Full",    CARD_MINT,     KICKER_GREEN),
        ("Bead + Fillering",      "Partial", CARD_PEACH,    ACCENT_AMBER),
        ("Tyre Building (6001)",  "Full",    CARD_MINT,     KICKER_GREEN),
        ("Curing (14811 fixed)",  "Out",     CARD_ROSE,     "B14444"),
    ]
    tbl = doc.add_table(rows=2, cols=len(stages))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _strip_table_borders(tbl)
    for col, (label, _, fill, accent) in enumerate(stages):
        cell = tbl.rows[0].cells[col]
        cell.width = Cm(17.5 / len(stages))
        _set_cell_shading(cell, fill)
        _set_cell_borders(cell, color="FFFFFF", size=12, sides=("left", "right"))
        _set_cell_margins(cell, top=120, left=80, bottom=120, right=80)
        p = _cell_first_para(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, label, bold=True, color=accent, size_pt=8.5)
    for col, (_, status, fill, accent) in enumerate(stages):
        cell = tbl.rows[1].cells[col]
        cell.width = Cm(17.5 / len(stages))
        _set_cell_shading(cell, fill)
        _set_cell_borders(cell, color="FFFFFF", size=12, sides=("left", "right"))
        _set_cell_margins(cell, top=60, left=80, bottom=120, right=80)
        p = _cell_first_para(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, status, color=SUBHEAD_GREY, size_pt=8.0)

    # Legend
    legend = doc.add_paragraph()
    legend.paragraph_format.space_before = Pt(4)
    legend.paragraph_format.space_after = Pt(0)
    _run(legend, "■  ", color=KICKER_GREEN, size_pt=9.0)
    _run(legend, "Full coverage    ", color=SUBHEAD_GREY, size_pt=8.5)
    _run(legend, "■  ", color=ACCENT_AMBER, size_pt=9.0)
    _run(legend, "Partial coverage    ", color=SUBHEAD_GREY, size_pt=8.5)
    _run(legend, "■  ", color="B14444", size_pt=9.0)
    _run(legend, "Out of scope / fixed", color=SUBHEAD_GREY, size_pt=8.5)


def _benefit_cards(doc):
    benefits = [
        ("92.3%",        "OTIF — Building → Curing",
         "36 / 38 productive blocks on time on the reference pilot run (V1 forward pass).",
         CARD_BLUE, ACCENT_BLUE),
        ("100%",         "Audit & reproducibility",
         "Byte-identical re-run on dag.json + DataFrame-equal on the 11-sheet workbook; 215+ tests.",
         CARD_MINT, KICKER_GREEN),
        ("0 silent",     "Defaults — every irregularity logged",
         "1 HALT + ~12 WARN findings surfaced from real input data; no value silently imputed.",
         CARD_LAVENDER, ACCENT_PURPLE),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _strip_table_borders(tbl)
    for col, (big, sub, body, fill, accent) in enumerate(benefits):
        cell = tbl.rows[0].cells[col]
        cell.width = Cm(17.5 / 3)
        _set_cell_shading(cell, fill)
        _set_cell_borders(cell, color="FFFFFF", size=14, sides=("left", "right"))
        _set_cell_margins(cell, top=220, left=200, bottom=220, right=200)
        p = _cell_first_para(cell)
        _run(p, big, bold=True, color=accent, size_pt=22.0)
        p.paragraph_format.space_after = Pt(2)
        p2 = cell.add_paragraph()
        _run(p2, sub, bold=True, color=HEADING_DARK, size_pt=9.5)
        p2.paragraph_format.space_after = Pt(2)
        p3 = cell.add_paragraph()
        _run(p3, body, color=SUBHEAD_GREY, size_pt=8.5)
        p3.paragraph_format.space_after = Pt(0)


def _two_card_row(doc, left_payload, right_payload):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _strip_table_borders(tbl)
    for col, (heading, items, fill, accent) in enumerate((left_payload, right_payload)):
        cell = tbl.rows[0].cells[col]
        cell.width = Cm(17.5 / 2)
        _set_cell_shading(cell, fill)
        _set_cell_borders(cell, color="FFFFFF", size=14, sides=("left", "right"))
        _set_cell_margins(cell, top=180, left=220, bottom=180, right=220)
        p = _cell_first_para(cell)
        _run(p, heading, bold=True, color=accent, size_pt=10.5)
        p.paragraph_format.space_after = Pt(4)
        for it in items:
            bp = cell.add_paragraph()
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(3)
            bp.paragraph_format.left_indent = Cm(0.3)
            _run(bp, "•  ", bold=True, color=accent, size_pt=9.0)
            _markdown_run(bp, it, base_color=HEADING_DARK, size_pt=9.0)


# -----------------------------------------------------------------------------
# Story
# -----------------------------------------------------------------------------
def build_docx(out_path: Path) -> Path:
    doc = Document()

    # Default styles → Arial
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    _set_margins(section)
    _set_footer(section)

    # ---- Title block ---------------------------------------------------------
    _two_col_title_table(doc)

    # ---- 01 — Problem & Solution --------------------------------------------
    _section_kicker(doc, "01 — Problem & Solution", color=ACCENT_BLUE)

    _card(doc, CARD_BLUE, ACCENT_BLUE,
          "Problem Statement",
          [
              "**Aging windows on every BOM edge** — calendered ply unsafe to build with after 24–48 h; master compounds behave differently outside their window.",
              "**7-level deep BOM**, eight mandatory components AND-joined at Tyre Building — spreadsheets cannot enforce this consistently.",
              "**Curing press 14811 is the fixed bottleneck** — once published, the plan is immovable; unfed blocks cascade across the 14-day horizon.",
              "**No deterministic audit trail** today — re-runs drift, infeasibilities are silently fixed, planners can’t reproduce yesterday’s decisions.",
          ])

    _card(doc, CARD_MINT, KICKER_GREEN,
          "Solution",
          [
              "**Forward, single-pass, BOM-driven scheduler** — walks BOM downward from each curing block, sizes lots with MPQ + aging-aware aggregation, dispatches in topological order with a CPM backward pass for per-lot floor/ceiling.",
              "**Hard constraints enforced, never silently fixed** — aging-MIN/MAX inclusive (L22), MPQ split, machine eligibility, atomic AND-join across 8 components, FEFO with soft reservation (L16/L19).",
              "**Flag-and-continue (L11)** — curing is never shifted; infeasibilities and aging breaches are logged with the binding constraint named, the engine moves on.",
              "**Reproducible by construction** — single integer-minute domain (L20), single ceil rounding, every sort explicit; byte-identical re-runs verified by integration tests.",
          ])

    # ---- 02 — Value Chain Coverage ------------------------------------------
    _section_kicker(doc, "02 — Value Chain Coverage", color=ACCENT_BLUE)
    _value_chain(doc)

    # ---- 03 — Quantifiable Benefits -----------------------------------------
    _section_kicker(doc, "03 — Quantifiable Benefits", color=ACCENT_BLUE)
    _benefit_cards(doc)

    add_p = doc.add_paragraph()
    add_p.paragraph_format.space_before = Pt(8)
    add_p.paragraph_format.space_after = Pt(2)
    _run(add_p, "Additional advantages", bold=True, color=HEADING_DARK, size_pt=10.0)
    advantages = [
        "**Auto-t₀ anchor (L17)** with guardrail HALT — refuses to commit lots if the longest BOM aging chain cannot reach the first curing block.",
        "**Eleven-sheet bundled workbook** — schedule, machine view, building-to-curing, aging violations, infeasibilities, reservation log, KPIs, audit findings — diff-friendly for planner review.",
        "**Stable HALT exit codes (10, 11, 12, 20, 30)** — CI / wrapper scripts branch on the binding finding without scraping logs.",
        "**Modular pipeline of 13 routes** — each step runs in isolation given the upstream typed result; debugging is incremental, not all-or-nothing.",
        "**Capstrip subtree (L12) auto-excluded** — out-of-scope BOM branch tagged in the visualisation; engine never tries to schedule against incomplete data.",
    ]
    for a in advantages:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.left_indent = Cm(0.35)
        _run(bp, "•  ", bold=True, color=ACCENT_BLUE, size_pt=9.5)
        _markdown_run(bp, a, base_color=HEADING_DARK, size_pt=9.5)

    # ---- 04 — Who benefits & Tech stack -------------------------------------
    _section_kicker(doc, "04 — Who Benefits & Tech Stack", color=ACCENT_BLUE)
    _two_card_row(
        doc,
        ("Primary beneficiaries",
         [
             "**Plant Planning Team** — owns the May curing plan and the upstream feasibility check.",
             "**Production Scheduling & Floor Supervisors** — execute the schedule by machine view.",
             "**Industrial Engineering** — reads aging violations + infeasibilities to spot bottlenecks.",
             "**Data Science / Analytics** — extends to V2 (changeover, utilisation, event-heap dispatch).",
         ],
         CARD_CREAM, HEADING_DARK),
        ("Tech stack / tools",
         [
             "Python ≥ 3.11   |   uv package manager   |   pytest + pytest-cov",
             "pandas   |   networkx   |   openpyxl",
             "matplotlib   |   plotly   |   graphviz   |   PyYAML",
             "**Spec discipline**: CLAUDE.md (23 locked decisions, 26-step flow, 5 golden fixtures).",
         ],
         CARD_LAVENDER, ACCENT_PURPLE),
    )

    # ---- 05 — Assumptions, Constraints & Next Steps -------------------------
    _section_kicker(doc, "05 — Assumptions, Constraints & Next Steps", color=ACCENT_BLUE)
    _two_card_row(
        doc,
        ("Assumptions & Constraints",
         [
             "Curing plan is immovable (L4.5) — engine never reschedules curing.",
             "Aging Master is authoritative; Buffer Master ignored (L4).",
             "Zero starting GT inventory; raws assumed bottomless pre-horizon (L2).",
             "Capstrip chain on ice (L12) — corrected Aging / MPQ data pending.",
             "**V1 changeover = 0 min** (L8); 15-min different-product model deferred to V2.",
             "Uniform 95 % efficiency factor across all operations (L10).",
             "Topological greedy dispatch in V1; strict event-heap (L21) deferred to V2.",
         ],
         CARD_PEACH, ACCENT_AMBER),
        ("Next Steps / Ask",
         [
             "**Planner sign-off** on the OTIF 92.3 % reference run and the 4 logged aging breaches.",
             "**Supply `proc_time` for `BD-12843443-4` Fillering** — current HALT (exit 10) is by design (§8.D).",
             "**Corrected Capstrip Aging Master + MPQ** so L12 can be lifted in V2.",
             "**V2 scope sign-off**: changeover modelling, LSF event dispatch, utilisation objective, same-product clustering on FRC + mixers.",
             "**Peer review** by plant engineering lead before the floor pilot.",
             "**Pilot the floor view** — extend Gantt from 3 sample blocks to all 42 if planner needs it.",
         ],
         CARD_ROSE, "B14444"),
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# -----------------------------------------------------------------------------
# PDF rendition (mirrors the same layout via reportlab)
# -----------------------------------------------------------------------------
def build_pdf(out_path: Path) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer, Table, TableStyle,
    )

    HEX = lambda h: colors.HexColor("#" + h)
    KGREEN, HDARK, SGREY, LGREY = HEX(KICKER_GREEN), HEX(HEADING_DARK), HEX(SUBHEAD_GREY), HEX(LABEL_GREY)
    BLUE, PURPLE, AMBER, BORDER = HEX(ACCENT_BLUE), HEX(ACCENT_PURPLE), HEX(ACCENT_AMBER), HEX(BORDER_GREY)
    F_BLUE, F_MINT, F_LAV, F_CREAM, F_PEACH, F_ROSE = (
        HEX(CARD_BLUE), HEX(CARD_MINT), HEX(CARD_LAVENDER),
        HEX(CARD_CREAM), HEX(CARD_PEACH), HEX(CARD_ROSE),
    )
    RED_TXT = HEX("B14444")

    # Paragraph XML expects '#RRGGBB'.
    def cx(rl_color):
        return "#" + format(rl_color.rgb() if hasattr(rl_color, "rgb") else 0, "06X") \
            if False else "#" + rl_color.hexval()[2:].upper()

    base = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=base["BodyText"], fontName="Helvetica",
                          fontSize=9.5, leading=12.5, textColor=HDARK)
    body_small = ParagraphStyle("bs", parent=body, fontSize=8.5, leading=11)

    def kicker_par(text, color=KGREEN, size=8):
        return Paragraph(
            f"<font name='Helvetica-Bold' size='{size}' color='#{color.hexval()[2:]}'>{text.upper()}</font>",
            ParagraphStyle("k", parent=body, spaceAfter=2, spaceBefore=0),
        )

    def title_par(text, size=20, color=HDARK):
        return Paragraph(
            f"<font name='Helvetica-Bold' size='{size}' color='#{color.hexval()[2:]}'>{text}</font>",
            ParagraphStyle("t", parent=body, spaceAfter=2, leading=size + 2),
        )

    def subtitle_par(text, color=SGREY, size=9.5):
        return Paragraph(
            f"<font name='Helvetica-Oblique' size='{size}' color='#{color.hexval()[2:]}'>{text}</font>",
            ParagraphStyle("st", parent=body, spaceAfter=0, leading=size + 2.5),
        )

    def label_value_par(label, value, value_color=HDARK):
        return Paragraph(
            f"<font name='Helvetica' size='8.5' color='#{LGREY.hexval()[2:]}'>{label}  </font>"
            f"<font name='Helvetica-Bold' size='8.5' color='#{value_color.hexval()[2:]}'>{value}</font>",
            ParagraphStyle("lv", parent=body, spaceAfter=2, leading=11),
        )

    def section_kicker(text, color=BLUE):
        return Paragraph(
            f"<font name='Helvetica-Bold' size='8.5' color='#{color.hexval()[2:]}'>{text.upper()}</font>",
            ParagraphStyle("sk", parent=body, spaceAfter=4, spaceBefore=10),
        )

    def md_to_xml(text):
        import re
        out = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        return out

    def bullet(text, color=BLUE, size=9, base_color=HDARK):
        return Paragraph(
            f"<font name='Helvetica-Bold' size='{size}' color='#{color.hexval()[2:]}'>•  </font>"
            f"<font name='Helvetica' size='{size}' color='#{base_color.hexval()[2:]}'>{md_to_xml(text)}</font>",
            ParagraphStyle("bu", parent=body, leftIndent=10, spaceAfter=3, leading=size + 3),
        )

    def card(fill, accent, title_text, items, base_color=HDARK, item_size=9):
        rows = [[title_par(title_text, size=11, color=accent)]]
        for it in items:
            rows.append([bullet(it, color=accent, size=item_size, base_color=base_color)])
        t = Table(rows, colWidths=[17.5 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), fill),
            ("LEFTPADDING", (0, 0), (-1, -1), 12),
            ("RIGHTPADDING", (0, 0), (-1, -1), 12),
            ("TOPPADDING", (0, 0), (0, 0), 10),
            ("BOTTOMPADDING", (0, -1), (-1, -1), 10),
            ("TOPPADDING", (0, 1), (-1, -1), 1),
            ("BOTTOMPADDING", (0, 1), (-1, -2), 1),
            ("BOX", (0, 0), (-1, -1), 0, fill),
        ]))
        return t

    def two_card_row(left, right):
        # left/right = (title, items, fill, accent)
        def render(payload):
            title, items, fill, accent = payload
            rows = [[title_par(title, size=10.5, color=accent)]]
            for it in items:
                rows.append([bullet(it, color=accent, size=9, base_color=HDARK)])
            t = Table(rows, colWidths=[8.4 * cm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), fill),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (0, 0), 10),
                ("BOTTOMPADDING", (0, -1), (-1, -1), 10),
                ("TOPPADDING", (0, 1), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 1), (-1, -2), 1),
            ]))
            return t

        outer = Table([[render(left), render(right)]], colWidths=[8.7 * cm, 8.7 * cm])
        outer.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (0, 0), 6),
            ("LEFTPADDING", (1, 0), (1, 0), 6),
            ("RIGHTPADDING", (1, 0), (-1, -1), 0),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return outer

    def benefit_cards():
        def card3(big, sub, body_text, fill, accent):
            rows = [
                [Paragraph(f"<font name='Helvetica-Bold' size='22' color='#{accent.hexval()[2:]}'>{big}</font>",
                           ParagraphStyle("big", parent=body, leading=24, spaceAfter=2))],
                [Paragraph(f"<font name='Helvetica-Bold' size='10' color='#{HDARK.hexval()[2:]}'>{sub}</font>",
                           ParagraphStyle("sub", parent=body, leading=12, spaceAfter=2))],
                [Paragraph(f"<font name='Helvetica' size='8.5' color='#{SGREY.hexval()[2:]}'>{body_text}</font>",
                           ParagraphStyle("bd", parent=body, leading=11))],
            ]
            t = Table(rows, colWidths=[5.55 * cm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), fill),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (0, 0), 14),
                ("BOTTOMPADDING", (0, -1), (-1, -1), 14),
            ]))
            return t

        b1 = card3("92.3%", "OTIF — Building → Curing",
                   "36 / 38 productive blocks on time on the reference pilot run (V1 forward pass).",
                   F_BLUE, BLUE)
        b2 = card3("100%", "Audit &amp; reproducibility",
                   "Byte-identical re-run on dag.json + DataFrame-equal on the 11-sheet workbook; 215+ tests.",
                   F_MINT, KGREEN)
        b3 = card3("0 silent", "Defaults — every irregularity logged",
                   "1 HALT + ~12 WARN findings surfaced from real input data; no value silently imputed.",
                   F_LAV, PURPLE)
        outer = Table([[b1, b2, b3]], colWidths=[5.85 * cm, 5.85 * cm, 5.85 * cm])
        outer.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (0, 0), 0),
            ("RIGHTPADDING", (0, 0), (0, 0), 4),
            ("LEFTPADDING", (1, 0), (1, 0), 4),
            ("RIGHTPADDING", (1, 0), (1, 0), 4),
            ("LEFTPADDING", (2, 0), (2, 0), 4),
            ("RIGHTPADDING", (2, 0), (2, 0), 0),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return outer

    def value_chain():
        stages = [
            ("Mixing",                "Full",    F_MINT,  KGREEN),
            ("Final Mixing",          "Full",    F_MINT,  KGREEN),
            ("Calendering (FRC)",     "Full",    F_MINT,  KGREEN),
            ("Belt + Ply Cutter",     "Full",    F_MINT,  KGREEN),
            ("Extrusion",             "Full",    F_MINT,  KGREEN),
            ("Bead + Fillering",      "Partial", F_PEACH, AMBER),
            ("Tyre Building (6001)",  "Full",    F_MINT,  KGREEN),
            ("Curing (14811 fixed)",  "Fixed",   F_ROSE,  RED_TXT),
        ]
        col_w = 17.5 / len(stages) * cm
        headers = []
        statuses = []
        for label, st, fill, accent in stages:
            headers.append(Paragraph(
                f"<para alignment='center'><font name='Helvetica-Bold' size='8.5' color='#{accent.hexval()[2:]}'>{label}</font></para>",
                ParagraphStyle("vh", parent=body, leading=10)))
            statuses.append(Paragraph(
                f"<para alignment='center'><font name='Helvetica' size='8' color='#{SGREY.hexval()[2:]}'>{st}</font></para>",
                ParagraphStyle("vs", parent=body, leading=10)))
        tbl = Table([headers, statuses], colWidths=[col_w] * len(stages))
        cmds = [
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]
        for col, (_, _, fill, _) in enumerate(stages):
            cmds.append(("BACKGROUND", (col, 0), (col, 1), fill))
            cmds.append(("LINEBEFORE", (col, 0), (col, 1), 4, colors.white))
            cmds.append(("LINEAFTER", (col, 0), (col, 1), 4, colors.white))
        tbl.setStyle(TableStyle(cmds))
        legend = Paragraph(
            f"<font name='Helvetica-Bold' color='#{KGREEN.hexval()[2:]}'>■</font> "
            f"<font name='Helvetica' size='8.5' color='#{SGREY.hexval()[2:]}'>Full coverage   </font>"
            f"<font name='Helvetica-Bold' color='#{AMBER.hexval()[2:]}'>■</font> "
            f"<font name='Helvetica' size='8.5' color='#{SGREY.hexval()[2:]}'>Partial coverage   </font>"
            f"<font name='Helvetica-Bold' color='#{RED_TXT.hexval()[2:]}'>■</font> "
            f"<font name='Helvetica' size='8.5' color='#{SGREY.hexval()[2:]}'>Out of scope / fixed</font>",
            ParagraphStyle("lg", parent=body, spaceBefore=4, spaceAfter=0, leading=11),
        )
        return [tbl, Spacer(1, 1 * mm), legend]

    def header_block():
        left_cell = [
            kicker_par("FORWARD PRODUCTION SCHEDULER", color=KGREEN, size=8.5),
            title_par("JK Tyre BTP — Passenger-Car Radial Pilot", size=19, color=HDARK),
            subtitle_par(
                "A deterministic, BOM-driven single-pass scheduler that converts a fixed May "
                "curing plan into a fully time-stamped, machine-assigned upstream schedule "
                "across the eight in-scope PCR components."),
        ]
        right_cell = [
            label_value_par("Author", "Anmol Saini"),
            label_value_par("Role", "Data Science Intern"),
            label_value_par("Date", "24 May 2026"),
            label_value_par("Team", "Data Science Team"),
            Paragraph(
                f"<font name='Helvetica' size='8.5' color='#{LGREY.hexval()[2:]}'>Readiness</font>",
                ParagraphStyle("rd", parent=body, spaceAfter=2, leading=11)),
            Paragraph(
                f"<font name='Helvetica' size='8.5' color='#{LGREY.hexval()[2:]}'>○ Concept   </font>"
                f"<font name='Helvetica-Bold' size='8.5' color='#{BLUE.hexval()[2:]}'>● Prototype</font>"
                f"<font name='Helvetica' size='8.5' color='#{LGREY.hexval()[2:]}'>   ○ Production</font>",
                ParagraphStyle("rd2", parent=body, spaceAfter=2, leading=11)),
            label_value_par("Scope", "V1 — Demand Fulfilment"),
        ]
        tbl = Table([[left_cell, right_cell]], colWidths=[12.3 * cm, 5.2 * cm])
        tbl.setStyle(TableStyle([
            ("LINEBEFORE", (1, 0), (1, 0), 0.6, BORDER),
            ("LEFTPADDING", (0, 0), (0, 0), 0),
            ("RIGHTPADDING", (0, 0), (0, 0), 8),
            ("LEFTPADDING", (1, 0), (1, 0), 12),
            ("RIGHTPADDING", (1, 0), (1, 0), 0),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return tbl

    story = []
    story.append(header_block())
    story.append(Spacer(1, 3 * mm))

    # 01 — Problem & Solution
    story.append(section_kicker("01 — Problem & Solution"))
    story.append(card(F_BLUE, BLUE, "Problem Statement", [
        "**Aging windows on every BOM edge** — calendered ply unsafe to build with after 24–48 h; master compounds behave differently outside their window.",
        "**7-level deep BOM**, eight mandatory components AND-joined at Tyre Building — spreadsheets cannot enforce this consistently.",
        "**Curing press 14811 is the fixed bottleneck** — once published, the plan is immovable; unfed blocks cascade across the 14-day horizon.",
        "**No deterministic audit trail** today — re-runs drift, infeasibilities are silently fixed, planners can’t reproduce yesterday’s decisions.",
    ]))
    story.append(Spacer(1, 2 * mm))
    story.append(card(F_MINT, KGREEN, "Solution", [
        "**Forward, single-pass, BOM-driven scheduler** — walks BOM downward from each curing block, sizes lots with MPQ + aging-aware aggregation, dispatches in topological order with a CPM backward pass.",
        "**Hard constraints enforced, never silently fixed** — aging-MIN/MAX inclusive (L22), MPQ split, machine eligibility, atomic AND-join across 8 components, FEFO with soft reservation (L16/L19).",
        "**Flag-and-continue (L11)** — curing is never shifted; infeasibilities + aging breaches logged with binding constraint named.",
        "**Reproducible by construction** — single integer-minute domain (L20), single ceil rounding; byte-identical re-runs verified by integration tests.",
    ]))

    # 02 — Value chain
    story.append(section_kicker("02 — Value Chain Coverage"))
    for item in value_chain():
        story.append(item)

    # 03 — Quantifiable benefits
    story.append(section_kicker("03 — Quantifiable Benefits"))
    story.append(benefit_cards())
    story.append(Spacer(1, 3 * mm))
    story.append(Paragraph(
        f"<font name='Helvetica-Bold' size='10' color='#{HDARK.hexval()[2:]}'>Additional advantages</font>",
        ParagraphStyle("aa", parent=body, spaceBefore=2, spaceAfter=2)))
    for adv in [
        "**Auto-t₀ anchor (L17)** with guardrail HALT — refuses to commit lots if the longest BOM aging chain cannot reach the first curing block.",
        "**Eleven-sheet bundled workbook** — schedule, machine view, building-to-curing, aging violations, infeasibilities, reservation log, KPIs, audit findings — diff-friendly for planner review.",
        "**Stable HALT exit codes (10, 11, 12, 20, 30)** — CI / wrapper scripts branch on the binding finding without scraping logs.",
        "**Modular pipeline of 13 routes** — each step runs in isolation given the upstream typed result; debugging is incremental, not all-or-nothing.",
        "**Capstrip subtree (L12) auto-excluded** — out-of-scope BOM branch tagged in the visualisation; engine never tries to schedule against incomplete data.",
    ]:
        story.append(bullet(adv, color=BLUE, size=9, base_color=HDARK))

    # 04 — Who benefits & Tech stack
    story.append(section_kicker("04 — Who Benefits & Tech Stack"))
    story.append(two_card_row(
        ("Primary beneficiaries", [
            "**Plant Planning Team** — owns the May curing plan and the upstream feasibility check.",
            "**Production Scheduling & Floor Supervisors** — execute the schedule by machine view.",
            "**Industrial Engineering** — reads aging violations + infeasibilities to spot bottlenecks.",
            "**Data Science / Analytics** — extends to V2 (changeover, utilisation, event-heap dispatch).",
        ], F_CREAM, HDARK),
        ("Tech stack / tools", [
            "Python ≥ 3.11   |   uv package manager   |   pytest + pytest-cov",
            "pandas   |   networkx   |   openpyxl",
            "matplotlib   |   plotly   |   graphviz   |   PyYAML",
            "**Spec discipline**: CLAUDE.md (23 locked decisions, 26-step flow, 5 golden fixtures).",
        ], F_LAV, PURPLE),
    ))

    # 05 — Assumptions, constraints & next steps
    story.append(section_kicker("05 — Assumptions, Constraints & Next Steps"))
    story.append(two_card_row(
        ("Assumptions & Constraints", [
            "Curing plan is immovable (L4.5) — engine never reschedules curing.",
            "Aging Master is authoritative; Buffer Master ignored (L4).",
            "Zero starting GT inventory; raws assumed bottomless pre-horizon (L2).",
            "Capstrip chain on ice (L12) — corrected Aging / MPQ data pending.",
            "**V1 changeover = 0 min** (L8); 15-min different-product model deferred to V2.",
            "Uniform 95 % efficiency factor across all operations (L10).",
            "Topological greedy dispatch in V1; strict event-heap (L21) deferred to V2.",
        ], F_PEACH, AMBER),
        ("Next Steps / Ask", [
            "**Planner sign-off** on the OTIF 92.3 % reference run and the 4 logged aging breaches.",
            "**Supply `proc_time` for `BD-12843443-4` Fillering** — current HALT (exit 10) is by design (§8.D).",
            "**Corrected Capstrip Aging Master + MPQ** so L12 can be lifted in V2.",
            "**V2 scope sign-off**: changeover modelling, LSF event dispatch, utilisation objective, same-product clustering.",
            "**Peer review** by plant engineering lead before the floor pilot.",
            "**Pilot the floor view** — extend Gantt from 3 sample blocks to all 42 if planner needs it.",
        ], F_ROSE, RED_TXT),
    ))

    # Build doc
    def _footer(canv, doc):
        canv.saveState()
        canv.setFillColor(LGREY)
        canv.setFont("Helvetica", 8)
        canv.drawString(1.5 * cm, 0.9 * cm,
                        "CONFIDENTIAL — Internal Use Only          Page %d" % doc.page)
        canv.restoreState()

    doc = BaseDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=1.4 * cm, bottomMargin=1.4 * cm,
        title="JK Tyre BTP — Forward Production Scheduler (V1) — Proforma",
        author="Anmol Saini",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height,
                  id="body", showBoundary=0)
    doc.addPageTemplates([PageTemplate(id="Body", frames=[frame], onPage=_footer)])
    doc.build(story)
    return out_path


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    here = Path(__file__).resolve().parent
    docx_path = build_docx(here / "BTP_V1_Proforma_AnmolSaini.docx")
    print(f"Wrote {docx_path}")
    pdf_path = build_pdf(here / "BTP_V1_Proforma_AnmolSaini.pdf")
    print(f"Wrote {pdf_path}")
